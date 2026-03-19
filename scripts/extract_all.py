#!/usr/bin/env python3
"""
批量提取所有指南文本
遍历 guidelines/ 目录，提取所有 PDF 和 DOCX 文件
"""

import argparse
import sys
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor, as_completed


def extract_file(file_path: Path, force: bool = False) -> tuple[str, bool, str]:
    """提取单个文件"""
    output_dir = file_path.parent / "extracted"
    output_file = output_dir / f"{file_path.stem}.txt"

    if output_file.exists() and not force:
        return file_path.name, True, "已存在（跳过）"

    output_dir.mkdir(parents=True, exist_ok=True)

    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        import subprocess

        try:
            result = subprocess.run(
                ["pdftotext", "-layout", str(file_path), str(output_file)],
                capture_output=True,
                text=True,
            )
            if result.returncode == 0:
                return file_path.name, True, "✓ 提取完成"
            else:
                return file_path.name, False, f"pdftotext 错误: {result.stderr[:100]}"
        except Exception as e:
            return file_path.name, False, str(e)

    elif suffix == ".docx":
        try:
            from docx import Document

            doc = Document(file_path)
            lines = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    style_name = para.style.name if para.style else ""
                    if "Heading 1" in style_name or "标题 1" in style_name:
                        lines.append(f"# {text}\n")
                    elif "Heading 2" in style_name or "标题 2" in style_name:
                        lines.append(f"## {text}\n")
                    elif "Heading 3" in style_name or "标题 3" in style_name:
                        lines.append(f"### {text}\n")
                    else:
                        lines.append(text)

            if doc.tables:
                lines.append("\n\n--- Tables ---\n")
                for i, table in enumerate(doc.tables, 1):
                    lines.append(f"\n### Table {i}\n")
                    for row in table.rows:
                        cells = [
                            cell.text.strip().replace("\n", " ") for cell in row.cells
                        ]
                        lines.append("| " + " | ".join(cells) + " |")

            output_file.write_text("\n".join(lines), encoding="utf-8")
            return file_path.name, True, "✓ 提取完成"
        except Exception as e:
            return file_path.name, False, str(e)

    else:
        return file_path.name, False, f"不支持的文件类型: {suffix}"


def main():
    parser = argparse.ArgumentParser(description="批量提取指南文本")
    parser.add_argument(
        "--force",
        action="store_true",
        help="强制重新提取（覆盖已有文件）",
    )
    parser.add_argument(
        "--guidelines-dir",
        type=Path,
        default=Path("guidelines"),
        help="指南根目录（默认为 guidelines/）",
    )
    args = parser.parse_args()

    guidelines_dir = args.guidelines_dir.resolve()
    if not guidelines_dir.exists():
        print(f"指南目录不存在: {guidelines_dir}", file=sys.stderr)
        sys.exit(1)

    # 收集所有 PDF 和 DOCX 文件
    files = []
    for subdir in guidelines_dir.iterdir():
        if subdir.is_dir() and subdir.name != "extracted":
            files.extend(subdir.glob("*.pdf"))
            files.extend(subdir.glob("*.PDF"))
            files.extend(subdir.glob("*.docx"))
            files.extend(subdir.glob("*.DOCX"))

    if not files:
        print("未找到任何 PDF 或 DOCX 文件")
        sys.exit(0)

    print(f"发现 {len(files)} 个文件需要处理\n")

    success = 0
    failed = 0
    skipped = 0

    for file_path in sorted(files):
        name, ok, msg = extract_file(file_path, force=args.force)
        status = "✓" if ok else "✗"
        print(f"  {status} {name}: {msg}")
        if ok:
            if "跳过" in msg:
                skipped += 1
            else:
                success += 1
        else:
            failed += 1

    print(f"\n统计: 成功 {success}, 跳过 {skipped}, 失败 {failed}")


if __name__ == "__main__":
    main()
