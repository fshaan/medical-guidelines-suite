#!/usr/bin/env python3
"""
DOCX 文本提取脚本
使用 python-docx 提取正文段落和表格
"""

import argparse
import sys
from pathlib import Path


def extract_docx(docx_path: Path, output_path: Path) -> bool:
    """使用 python-docx 提取内容"""
    try:
        from docx import Document

        doc = Document(docx_path)
        lines = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # 根据样式判断标题层级
            style_name = para.style.name if para.style else ""
            if "Heading 1" in style_name or "标题 1" in style_name:
                lines.append(f"# {text}\n")
            elif "Heading 2" in style_name or "标题 2" in style_name:
                lines.append(f"## {text}\n")
            elif "Heading 3" in style_name or "标题 3" in style_name:
                lines.append(f"### {text}\n")
            else:
                lines.append(text)

        # 提取表格
        if doc.tables:
            lines.append("\n\n--- Tables ---\n")
            for i, table in enumerate(doc.tables, 1):
                lines.append(f"\n### Table {i}\n")
                first_row = True
                for row in table.rows:
                    cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                    lines.append("| " + " | ".join(cells) + " |")
                    # 第一行后添加分隔线
                    if first_row:
                        lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
                        first_row = False

        output_path.write_text("\n".join(lines), encoding="utf-8")
        return True

    except ImportError:
        print("python-docx 未安装", file=sys.stderr)
        return False
    except Exception as e:
        print(f"提取失败: {e}", file=sys.stderr)
        return False


def main():
    parser = argparse.ArgumentParser(description="提取 DOCX 文本")
    parser.add_argument("input", type=Path, help="输入 DOCX 文件路径")
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=None,
        help="输出目录（默认为 DOCX 同目录的 extracted/ 子目录）",
    )
    args = parser.parse_args()

    docx_path = args.input.resolve()
    if not docx_path.exists():
        print(f"文件不存在: {docx_path}", file=sys.stderr)
        sys.exit(1)

    # 确定输出目录
    if args.output_dir:
        output_dir = args.output_dir
    else:
        output_dir = docx_path.parent / "extracted"
    output_dir.mkdir(parents=True, exist_ok=True)

    # 提取内容
    base_name = docx_path.stem
    text_output = output_dir / f"{base_name}.txt"
    print(f"提取: {docx_path.name} -> {text_output.name}")
    if extract_docx(docx_path, text_output):
        print(f"✓ 提取完成: {text_output}")
    else:
        sys.exit(1)


if __name__ == "__main__":
    main()
